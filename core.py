"""
Lógica pura del análisis de predio.
Se puede importar desde cualquier automatización sin arrastrar la UI.
"""
import os, zipfile, io, calendar, base64, logging
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional, Tuple

import requests
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from shapely.geometry import Polygon, MultiPolygon
import geopandas as gpd

logger = logging.getLogger(__name__)

CARTOGRAFIA_DIR  = Path(os.getenv("CARTOGRAFIA_DIR", "./data"))
UC_SHP_NAME      = os.getenv("UC_SHAPEFILE_NAME",   "Unidades_Cronoestratigraficas.shp")
EXCEL_NAME       = os.getenv("GEOFORMA_EXCEL_NAME", "Diccionario_UC_Geomorfologia.xlsx")
UC_SHP_PATH      = CARTOGRAFIA_DIR / UC_SHP_NAME
EXCEL_PATH       = CARTOGRAFIA_DIR / EXCEL_NAME

MONTH_LABELS = ["Ene","Feb","Mar","Abr","May","Jun","Jul","Ago","Sep","Oct","Nov","Dic"]


def parse_kmz(file_bytes: bytes):
    """Extrae el primer polígono de un KMZ. Devuelve geometría Shapely."""
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        kml_name = next((f for f in z.namelist() if f.lower().endswith(".kml")), None)
        if not kml_name:
            raise ValueError("El KMZ no contiene un archivo KML válido.")
        kml_bytes = z.read(kml_name)

    root = ET.fromstring(kml_bytes)
    ns = "http://www.opengis.net/kml/2.2"
    polygons = []

    for coords_el in root.iter(f"{{{ns}}}coordinates"):
        text = coords_el.text.strip()
        points = []
        for triple in text.split():
            parts = triple.split(",")
            if len(parts) >= 2:
                try:
                    lon, lat = float(parts[0]), float(parts[1])
                    points.append((lon, lat))
                except ValueError:
                    continue
        if len(points) >= 3:
            polygons.append(Polygon(points))

    if not polygons:
        raise ValueError("No se encontró ningún polígono en el KMZ.")
    return polygons[0] if len(polygons) == 1 else MultiPolygon(polygons)


def fetch_power(lat: float, lon: float, start: int = 2015, end: int = 2025) -> pd.DataFrame:
    """Descarga serie mensual de NASA POWER (sin API key)."""
    url = "https://power.larc.nasa.gov/api/temporal/monthly/point"
    r = requests.get(url, params={
        "parameters": "PRECTOTCORR,T2M,T2M_MAX,T2M_MIN",
        "community": "AG",
        "longitude": round(lon, 4),
        "latitude":  round(lat, 4),
        "start": start, "end": end, "format": "JSON",
    }, timeout=60)
    r.raise_for_status()
    data = r.json()["properties"]["parameter"]

    rows = []
    for year in range(start, end + 1):
        for month in range(1, 13):
            key = f"{year}{month:02d}"
            pd_ = data["PRECTOTCORR"].get(key, -999)
            t2m = data["T2M"].get(key, -999)
            if pd_ == -999 or t2m == -999:
                continue
            rows.append({
                "year": year, "month": month,
                "prec_mm": round(pd_ * calendar.monthrange(year, month)[1], 1),
                "t2m":   t2m,
                "t_max": data["T2M_MAX"].get(key, t2m),
                "t_min": data["T2M_MIN"].get(key, t2m),
            })
    if not rows:
        raise ValueError("NASA POWER no devolvió datos para estas coordenadas.")
    return pd.DataFrame(rows)


def calc_stats(df: pd.DataFrame) -> dict:
    m = df.groupby("month").agg(
        prec=("prec_mm", "mean"), t2m=("t2m", "mean"),
        t_max=("t_max", "mean"), t_min=("t_min", "mean"),
    ).reset_index()

    m["seco"] = m["prec"] < 2 * m["t2m"]
    prec = m["prec"].tolist()
    peaks = sum(1 for i in range(12) if prec[i] > prec[(i-1)%12] and prec[i] > prec[(i+1)%12])

    return {
        "prec_anual":   round(m["prec"].sum(), 0),
        "t_media":      round(m["t2m"].mean(), 1),
        "t_max":        round(m["t_max"].max(), 1),
        "t_min":        round(m["t_min"].min(), 1),
        "prec_max_mes": round(m["prec"].max(), 0),
        "n_secos":      int(m["seco"].sum()),
        "regimen":      "bimodal" if peaks >= 2 else "monomodal",
        "monthly":      m.to_dict(orient="records"),
    }


# ── Holdridge ─────────────────────────────────────────────────────────────────

def calc_biotemperatura(monthly_data: list) -> float:
    """Biotemperatura media anual (°C) según Holdridge.
    Aproximación con datos mensuales: trunca T2M mensual a [0, 30] y promedia.
    Para datos diarios sería más preciso, pero en zonas tropicales colombianas
    (T2M mensual siempre entre 0 y 30) el resultado coincide con la T media.
    """
    temps = [max(0.0, min(30.0, d["t2m"])) for d in monthly_data]
    return round(sum(temps) / len(temps), 2)


def clasificar_holdridge(biotemp: float, prec_anual: float) -> dict:
    """Clasifica zona de vida de Holdridge a partir de biotemperatura (°C)
    y precipitación anual (mm). Devuelve dict con código, nombre, piso,
    provincia de humedad y relación ETP/P.
    """
    # ETP anual ≈ biotemp * 58.93 (Holdridge)
    etp_anual = biotemp * 58.93
    etp_p = etp_anual / prec_anual if prec_anual > 0 else float("inf")

    # Piso altitudinal por biotemperatura
    if   biotemp >= 24:   piso, sufijo = "Tropical",     "T"
    elif biotemp >= 18:   piso, sufijo = "Premontano",   "PM"
    elif biotemp >= 12:   piso, sufijo = "Montano Bajo", "MB"
    elif biotemp >= 6:    piso, sufijo = "Montano",      "M"
    elif biotemp >= 3:    piso, sufijo = "Subalpino",    "SA"
    elif biotemp >= 1.5:  piso, sufijo = "Alpino",       "A"
    else:                 piso, sufijo = "Nival",        "N"

    # Provincia de humedad por relación ETP/P
    if   etp_p > 16:    provincia, formacion = "superárido",  "desierto"
    elif etp_p > 8:     provincia, formacion = "perárido",    "matorral desértico"
    elif etp_p > 4:     provincia, formacion = "árido",       "monte espinoso"
    elif etp_p > 2:     provincia, formacion = "semiárido",   "bosque muy seco"
    elif etp_p > 1:     provincia, formacion = "subhúmedo",   "bosque seco"
    elif etp_p > 0.5:   provincia, formacion = "húmedo",      "bosque húmedo"
    elif etp_p > 0.25:  provincia, formacion = "perhúmedo",   "bosque muy húmedo"
    else:               provincia, formacion = "superhúmedo", "bosque pluvial"

    # Código compacto: bs-T, bh-PM, etc.
    codigo_formacion = {
        "desierto": "d", "matorral desértico": "md", "monte espinoso": "me",
        "bosque muy seco": "bms", "bosque seco": "bs", "bosque húmedo": "bh",
        "bosque muy húmedo": "bmh", "bosque pluvial": "bp",
    }[formacion]
    codigo = f"{codigo_formacion}-{sufijo}"
    nombre = f"{formacion.capitalize()} {piso}"

    return {
        "codigo":    codigo,
        "nombre":    nombre,
        "piso":      piso,
        "provincia": provincia,
        "biotemp":   biotemp,
        "etp_anual": round(etp_anual, 0),
        "etp_p":     round(etp_p, 2),
    }


def texto_holdridge(holdridge: dict, prec_anual: float) -> str:
    """Genera el párrafo descriptivo de la zona de vida."""
    return (
        f"Según el sistema de zonas de vida de Holdridge, el predio se clasifica como "
        f"{holdridge['nombre']} ({holdridge['codigo']}). Esta clasificación se obtiene a partir "
        f"de una biotemperatura media anual de {holdridge['biotemp']:.1f}°C, una precipitación "
        f"total anual de {prec_anual:.0f} mm y una evapotranspiración potencial estimada de "
        f"{holdridge['etp_anual']:.0f} mm anuales (relación ETP/P = {holdridge['etp_p']:.2f}), "
        f"correspondiente a una provincia de humedad {holdridge['provincia']} dentro del piso "
        f"altitudinal {holdridge['piso']}.\n\n"
        f"Fuente: Holdridge, L. R. (1967). Life Zone Ecology. Tropical Science Center, San José, "
        f"Costa Rica. Cálculo basado en datos climáticos de NASA POWER. Elaboración propia."
    )


# ── Charts y mapas ────────────────────────────────────────────────────────────

def make_charts(monthly_data: list) -> Tuple[bytes, bytes]:
    """Devuelve (precip_png, temp_png) como bytes."""
    months = [d["month"] for d in monthly_data]
    prec   = [d["prec"]  for d in monthly_data]
    temp   = [d["t2m"]   for d in monthly_data]

    def _render(color, values, ylabel, title) -> bytes:
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(months, values, color=color, linewidth=2, marker="o", markersize=4)
        if color == "#1565C0":
            ax.fill_between(months, values, alpha=0.12, color=color)
        ax.set_xticks(range(1, 13)); ax.set_xticklabels(MONTH_LABELS)
        ax.set_ylabel(ylabel); ax.set_title(title); ax.grid(alpha=0.3)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
        plt.close(fig)
        return buf.getvalue()

    return (
        _render("#1565C0", prec, "Precipitación (mm/mes)", "Precipitación mensual promedio"),
        _render("#C62828", temp, "Temperatura media (°C)", "Temperatura mensual promedio"),
    )


def get_geoforma(
    predio_geom,
    uc_shp_path: Path = UC_SHP_PATH,
    excel_path:  Path = EXCEL_PATH,
    campo: str = "SimboloUC",
) -> Tuple[Optional[str], Optional[pd.DataFrame], Optional[object]]:
    """Intersecta el predio con el shapefile de UC y busca en el diccionario Excel.
    Retorna (simbolo, df_diccionario, geometria_uc_en_4326).
    """
    if not Path(uc_shp_path).exists():
        raise FileNotFoundError(f"Shapefile no encontrado: {uc_shp_path}")
    if not Path(excel_path).exists():
        raise FileNotFoundError(f"Excel no encontrado: {excel_path}")

    uc_gdf = gpd.read_file(str(uc_shp_path))
    if campo not in uc_gdf.columns:
        raise ValueError(f"El shapefile no tiene '{campo}'. Columnas: {list(uc_gdf.columns)}")

    predio_gdf = gpd.GeoDataFrame(geometry=[predio_geom], crs="EPSG:4326")
    predio_proj = predio_gdf.to_crs("EPSG:3116")
    uc_proj = uc_gdf.to_crs("EPSG:3116") if uc_gdf.crs else uc_gdf.set_crs("EPSG:3116")

    inter = gpd.overlay(predio_proj, uc_proj, how="intersection")
    if inter.empty:
        return None, None, None

    inter["_area"] = inter.geometry.area
    simbolo = inter.groupby(campo)["_area"].sum().idxmax()

    # Solo los polígonos de la UC dominante que intersectan el predio
    predio_geom_proj = predio_proj.geometry.iloc[0]
    uc_match = uc_proj[
        (uc_proj[campo].astype(str) == str(simbolo)) &
        (uc_proj.intersects(predio_geom_proj))
    ]
    uc_geom_4326 = uc_match.to_crs("EPSG:4326").geometry.unary_union

    try:
        df_dict = pd.read_excel(str(excel_path), engine="xlrd")
    except Exception:
        df_dict = pd.read_excel(str(excel_path), engine="openpyxl")

    df_uc = df_dict[df_dict[campo].astype(str) == str(simbolo)]
    return simbolo, (df_uc if not df_uc.empty else None), uc_geom_4326


def make_mapa_satelital(predio_geom, uc_geom=None) -> Optional[bytes]:
    """Genera PNG satelital con el polígono de la geoforma (UC) al estilo del informe de referencia."""
    try:
        import contextily as ctx

        fig, ax = plt.subplots(figsize=(10, 8))

        if uc_geom is not None:
            uc_gdf = gpd.GeoDataFrame(geometry=[uc_geom], crs="EPSG:4326")
            uc_web = uc_gdf.to_crs(epsg=3857)
            # Rosa semi-transparente igual al documento de referencia
            uc_web.plot(ax=ax, facecolor="#F4A0C0", edgecolor="#CC3366",
                        alpha=0.55, linewidth=1.5, zorder=2)
            # Zoom a la geoforma con 5 % de margen
            bounds = uc_web.total_bounds
            dx = (bounds[2] - bounds[0]) * 0.05
            dy = (bounds[3] - bounds[1]) * 0.05
        else:
            predio_gdf = gpd.GeoDataFrame(geometry=[predio_geom], crs="EPSG:4326")
            predio_web = predio_gdf.to_crs(epsg=3857)
            bounds = predio_web.total_bounds
            dx = max((bounds[2] - bounds[0]) * 0.3, 1000)
            dy = max((bounds[3] - bounds[1]) * 0.3, 1000)

        ax.set_xlim(bounds[0] - dx, bounds[2] + dx)
        ax.set_ylim(bounds[1] - dy, bounds[3] + dy)

        ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, zoom="auto")
        ax.set_axis_off()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
        plt.close(fig)
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"No se pudo generar imagen satelital: {type(e).__name__}: {e}")
        return None


# ── Textos ────────────────────────────────────────────────────────────────────

def texto_clima(stats: dict, start: int, end: int) -> str:
    t = stats['t_media']
    clima_termico = "cálido" if t > 24 else "templado" if t > 18 else "frío" if t > 12 else "muy frío/páramo"
    return (
        f"El clima del predio se clasifica predominantemente como {clima_termico}, "
        f"con una temperatura media mensual que varía entre {stats['t_min']:.1f}°C y {stats['t_max']:.1f}°C "
        f"(promedio general de {stats['t_media']:.1f}°C).\n\n"
        f"El régimen de precipitación es {stats['regimen']}, alcanzando un acumulado anual promedio de "
        f"{stats['prec_anual']:.0f} mm para el periodo {start}-{end}. Durante la época seca las lluvias "
        f"disminuyen significativamente, identificándose {stats['n_secos']} meses secos por año "
        f"(criterio de Gaussen, P < 2T). En la temporada húmeda, los aportes hídricos llegan a topes de "
        f"{stats['prec_max_mes']:.0f} mm mensuales.\n\n"
        f"Fuente: NASA Prediction Of Worldwide Energy Resources (POWER), {start}-{end}. Elaboración propia."
    )


def texto_geoforma(simbolo: str, df_uc: pd.DataFrame) -> str:
    r0 = df_uc.iloc[0]
    bullets = "\n".join(
        f"• {r.get('Geoforma','')}: {r.get('DescGeoforma','')}"
        for _, r in df_uc.iterrows()
        if pd.notna(r.get("Geoforma")) and str(r.get("Geoforma")).strip()
    )
    return (
        f"De acuerdo con la cartografía geomorfológica disponible, el predio se ubica dentro de la unidad "
        f"geomorfológica {simbolo}, clasificada como \"{r0.get('NombreUC','')}\". Esta unidad corresponde a "
        f"{r0.get('TipoUnidad','')} de edad {r0.get('Edad','')}, compuesta por {r0.get('Materiales','')}. "
        f"{r0.get('ProcesoFormacion','')}\n\n"
        f"La unidad {simbolo} comprende las siguientes geoformas principales:\n{bullets}"
    )


# ── DOCX consolidado ──────────────────────────────────────────────────────────

def generate_docx(
    clima_txt: str,
    geo_txt: str,
    precip_png: bytes,
    temp_png: bytes,
    mapa_png: Optional[bytes] = None,
    holdridge_txt: Optional[str] = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches
    import io as _io

    doc = Document()
    doc.add_heading('Informe de Aprovechamiento Forestal', 0)

    doc.add_heading('1. Análisis Climático', level=1)
    doc.add_paragraph(clima_txt)
    doc.add_picture(_io.BytesIO(precip_png), width=Inches(5.5))
    doc.add_picture(_io.BytesIO(temp_png), width=Inches(5.5))

    if holdridge_txt:
        doc.add_heading('2. Zona de Vida (Holdridge)', level=1)
        doc.add_paragraph(holdridge_txt)
        geo_heading = '3. Análisis Geomorfológico'
    else:
        geo_heading = '2. Análisis Geomorfológico'

    doc.add_heading(geo_heading, level=1)
    if mapa_png:
        doc.add_picture(_io.BytesIO(mapa_png), width=Inches(5.5))
    if geo_txt:
        doc.add_paragraph(geo_txt)
    else:
        doc.add_paragraph("No se encontró información geomorfológica para el predio solicitado.")

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
